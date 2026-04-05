#!/usr/bin/env -S .venv/bin/python3
"""Resume tailoring pipeline — from job posting to ATS-friendly DOCX."""

import argparse
import json
import logging
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).resolve().parent

MASTER_RESUME = SCRIPT_DIR / "master_resume.json"
PROMPT_FILE = SCRIPT_DIR / "prompts" / "tailor_prompt.txt"
AESTHETICS_PROMPT_FILE = SCRIPT_DIR / "prompts" / "aesthetics_prompt.txt"
CONTENT_CRITIC_PROMPT_FILE = SCRIPT_DIR / "prompts" / "content_critic_prompt.txt"

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Step 1 — Claude tailoring (content writer)
# ---------------------------------------------------------------------------

def _strip_fences(raw: str) -> str:
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"```\s*$", "", raw, flags=re.MULTILINE)
    return raw.strip()


def call_claude(
    job_description: str,
    master: dict,
    target_words: int = 460,
    feedback: str | None = None,
) -> dict:
    """Send master resume + job description to Claude Code CLI and return tailored JSON.

    target_words: approximate combined word budget for summary + bullets + project descriptions.
    feedback: optional guidance from the previous iteration (fit + aesthetics).
    """
    template = PROMPT_FILE.read_text()
    feedback_block = (
        f"**FEEDBACK FROM PREVIOUS ATTEMPT — APPLY THESE CHANGES:**\n{feedback}"
        if feedback
        else ""
    )
    system_prompt = template.replace("{TARGET_WORDS}", str(target_words)).replace(
        "{FEEDBACK_BLOCK}", feedback_block
    )

    user_message = (
        f"Job Posting:\n{job_description}\n\n"
        f"Master Resume:\n{json.dumps(master)}"
    )

    result = subprocess.run(
        [
            "claude", "-p",
            "--system-prompt", system_prompt,
            "--output-format", "text",
            "--model", "sonnet",
            "--max-turns", "1",
        ],
        input=user_message,
        capture_output=True,
        text=True,
        check=True,
    )
    return json.loads(_strip_fences(result.stdout))


# ---------------------------------------------------------------------------
# Step 1.5 — Fit measurement + aesthetics review
# ---------------------------------------------------------------------------

def docx_to_pdf(docx_path: Path) -> Path:
    """Convert a .docx to .pdf via LibreOffice headless. Returns PDF path."""
    if not shutil.which("soffice"):
        raise RuntimeError("soffice (LibreOffice) not found in PATH — required for fit measurement")
    # Use an isolated profile to avoid collisions with a running LibreOffice.
    profile = Path("/tmp") / f"lo_profile_{docx_path.stem}"
    subprocess.run(
        [
            "soffice",
            f"-env:UserInstallation=file://{profile}",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(docx_path.parent),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
    )
    return docx_path.with_suffix(".pdf")


def measure_fit(pdf_path: Path) -> dict:
    """Return {pages, fill_ratio, page1_png_path}.

    fill_ratio on a single page = bottom-of-lowest-block / usable height (0..1).
    When content overflows, fill_ratio > 1.0 and reflects total content volume:
    1 + (page-2-content / usable), capped at 2.0 for 2 pages of full content.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(pdf_path)
    pages = doc.page_count
    page_h = doc[0].rect.height
    top_margin = 0.5 * 72
    bottom_margin = 0.5 * 72
    usable = page_h - top_margin - bottom_margin

    # Page 1 fill
    blocks_p1 = doc[0].get_text("blocks")
    if blocks_p1:
        content_h_p1 = max(b[3] for b in blocks_p1) - top_margin
        fill_p1 = max(0.0, min(1.0, content_h_p1 / usable))
    else:
        fill_p1 = 0.0

    fill_ratio = fill_p1
    if pages > 1:
        # Add overflow signal from page 2.
        blocks_p2 = doc[1].get_text("blocks")
        if blocks_p2:
            content_h_p2 = max(b[3] for b in blocks_p2) - top_margin
            fill_p2 = max(0.0, min(1.0, content_h_p2 / usable))
            fill_ratio = 1.0 + fill_p2

    png_path = pdf_path.with_suffix(".page1.png")
    pix = doc[0].get_pixmap(dpi=110)
    pix.save(str(png_path))
    doc.close()

    return {
        "pages": pages,
        "fill_ratio": round(fill_ratio, 3),
        "page1_png_path": png_path,
    }


def review_content(job_description: str, tailored: dict, master: dict) -> dict:
    """Critic pass: does the tailored resume actually match what the JD asks for, and is it faithful to master?

    Returns a dict with keys: acceptable, fidelity_violations, coverage_gaps,
    irrelevant_content, positioning_notes. Neutral passthrough on failure so
    the loop doesn't break.
    """
    try:
        system_prompt = CONTENT_CRITIC_PROMPT_FILE.read_text()
        user_message = (
            f"Job Posting:\n{job_description}\n\n"
            f"Tailored Resume:\n{json.dumps(tailored)}\n\n"
            f"Master Resume:\n{json.dumps(master)}"
        )
        result = subprocess.run(
            [
                "claude", "-p",
                "--system-prompt", system_prompt,
                "--output-format", "text",
                "--model", "sonnet",
                "--max-turns", "1",
            ],
            input=user_message,
            capture_output=True,
            text=True,
            check=True,
            timeout=180,
        )
        return json.loads(_strip_fences(result.stdout))
    except Exception as e:
        log.warning(f"content critic failed, skipping: {e}")
        return {
            "acceptable": True,
            "fidelity_violations": [],
            "coverage_gaps": [],
            "irrelevant_content": [],
            "positioning_notes": [],
        }


def _format_content_feedback(critique: dict) -> list[str]:
    """Turn a content-critic dict into feedback lines for the writer."""
    parts = []
    violations = critique.get("fidelity_violations") or []
    if violations:
        parts.append("FIDELITY VIOLATIONS (must fix — these break the resume):")
        parts.extend(f"- {v}" for v in violations)
    gaps = critique.get("coverage_gaps") or []
    if gaps:
        parts.append("COVERAGE GAPS (JD requirements not reflected in tailored resume):")
        parts.extend(f"- {g}" for g in gaps)
    irrelevant = critique.get("irrelevant_content") or []
    if irrelevant:
        parts.append("IRRELEVANT CONTENT (wasted space — swap for something matching the JD):")
        parts.extend(f"- {i}" for i in irrelevant)
    positioning = critique.get("positioning_notes") or []
    if positioning:
        parts.append("POSITIONING:")
        parts.extend(f"- {p}" for p in positioning)
    return parts


def review_aesthetics(png_path: Path, tailored: dict) -> dict:
    """Ask Claude (Sonnet, vision) to review the rendered page. Returns parsed JSON feedback.

    On any failure, returns a neutral passthrough so the loop doesn't break.
    """
    try:
        system_prompt = AESTHETICS_PROMPT_FILE.read_text()
        # @<abs-path> is the claude CLI's inline file reference syntax — it attaches the image.
        user_message = (
            f"Rendered resume page 1 (look at this image): @{png_path}\n\n"
            f"Tailored resume JSON:\n{json.dumps(tailored)}"
        )
        result = subprocess.run(
            [
                "claude", "-p",
                "--system-prompt", system_prompt,
                "--output-format", "text",
                "--model", "sonnet",
                "--max-turns", "1",
            ],
            input=user_message,
            capture_output=True,
            text=True,
            check=True,
            timeout=180,
        )
        return json.loads(_strip_fences(result.stdout))
    except Exception as e:
        log.warning(f"aesthetics review failed, skipping: {e}")
        return {"acceptable": True, "issues": [], "suggestions": []}


def _score(fit: dict) -> float:
    """Higher is better. 1 page + high fill is ideal; overflow is heavily penalized."""
    if fit["pages"] > 1:
        # Stronger penalty the further we overflow; still ordered below any 1-page result.
        return -10.0 - (fit["fill_ratio"] - 1.0)
    return fit["fill_ratio"]


def panic_trim(tailored: dict, out_dir: Path, max_trims: int = 8) -> tuple[dict, dict]:
    """Deterministic last-resort: pop bullets from the tailored JSON until it fits on one page.

    Strategy: repeatedly drop the last bullet from the experience entry that has the most
    bullets (tie-break: last entry). Re-render + re-measure after each trim.
    """
    docx_path = out_dir / "resume.docx"
    for i in range(max_trims):
        # Find the entry with the most bullets.
        if not tailored.get("experience"):
            break
        idx_max = max(
            range(len(tailored["experience"])),
            key=lambda j: (len(tailored["experience"][j].get("bullets", [])), j),
        )
        bullets = tailored["experience"][idx_max].get("bullets", [])
        if len(bullets) <= 2:
            # Don't gut the entry. Try trimming a project instead.
            if tailored.get("projects") and len(tailored["projects"]) > 1:
                tailored["projects"].pop()
            else:
                break
        else:
            bullets.pop()

        render_docx(tailored, docx_path)
        pdf_path = docx_to_pdf(docx_path)
        fit = measure_fit(pdf_path)
        log.info(f"panic_trim {i}: pages={fit['pages']} fill={fit['fill_ratio']}")
        if fit["pages"] == 1:
            return tailored, fit
    return tailored, fit


def tailor_with_loop(
    job_description: str,
    master: dict,
    out_dir: Path,
    max_iterations: int = 3,
) -> tuple[dict, dict, list]:
    """Run writer → render → measure → review loop. Returns (best_tailored, best_fit, history)."""
    target_words = 400
    feedback: str | None = None
    history = []

    best_tailored = None
    best_fit = None
    best_score = -999.0
    docx_path = out_dir / "resume.docx"

    for i in range(max_iterations):
        log.info(f"iter {i}: target_words={target_words}")
        tailored = call_claude(job_description, master, target_words, feedback)
        render_docx(tailored, docx_path)
        pdf_path = docx_to_pdf(docx_path)
        fit = measure_fit(pdf_path)
        aesthetics = review_aesthetics(fit["page1_png_path"], tailored)
        content = review_content(job_description, tailored, master)

        score = _score(fit)
        content_ok = content.get("acceptable", True) and not content.get("fidelity_violations")
        log.info(
            f"iter {i}: pages={fit['pages']} fill={fit['fill_ratio']} "
            f"score={score:.3f} aesthetic_ok={aesthetics.get('acceptable')} "
            f"content_ok={content_ok}"
        )
        history.append(
            {
                "iteration": i,
                "target_words": target_words,
                "pages": fit["pages"],
                "fill_ratio": fit["fill_ratio"],
                "aesthetics": aesthetics,
                "content": content,
            }
        )

        # Score only rewards layout; content quality gates acceptance separately.
        if score > best_score:
            best_score = score
            best_tailored = tailored
            best_fit = fit

        # Accept: exactly 1 page, fill >= 0.92, both reviewers happy, no fidelity violations
        if (
            fit["pages"] == 1
            and fit["fill_ratio"] >= 0.92
            and aesthetics.get("acceptable", True)
            and content_ok
        ):
            log.info(f"iter {i}: accepted")
            break

        # Build feedback for next iteration
        parts = []
        if fit["pages"] > 1:
            parts.append(
                f"Previous attempt OVERFLOWED to {fit['pages']} pages. Cut content aggressively — "
                f"drop the weakest bullet from each experience entry and drop 1 project. "
                f"Do NOT merge entries to save space."
            )
            target_words = int(target_words * 0.78)
        elif fit["fill_ratio"] < 0.92:
            shortfall = 0.95 - fit["fill_ratio"]
            parts.append(
                f"Previous attempt only filled {int(fit['fill_ratio']*100)}% of the page. "
                f"Expand content to fill the page — add more bullets, projects, or detail."
            )
            target_words = min(int(target_words * (1 + shortfall * 1.5)), 650)

        content_parts = _format_content_feedback(content)
        if content_parts:
            parts.append("Content critic feedback:")
            parts.extend(content_parts)

        suggestions = aesthetics.get("suggestions", []) if isinstance(aesthetics, dict) else []
        if suggestions:
            parts.append("Aesthetic reviewer suggestions:")
            parts.extend(f"- {s}" for s in suggestions)

        feedback = "\n".join(parts) if parts else None

    # Re-render the best one to ensure resume.docx matches what we return
    if best_tailored is not None:
        render_docx(best_tailored, docx_path)
        # Panic trim: if still overflowing, deterministically drop bullets until it fits.
        if best_fit and best_fit["pages"] > 1:
            log.info("entering panic_trim — LLM loop couldn't fit on one page")
            best_tailored, panic_fit = panic_trim(best_tailored, out_dir)
            if panic_fit:
                best_fit = panic_fit
                history.append(
                    {
                        "iteration": "panic_trim",
                        "target_words": None,
                        "pages": panic_fit["pages"],
                        "fill_ratio": panic_fit["fill_ratio"],
                        "aesthetics": {"acceptable": True, "issues": [], "suggestions": []},
                    }
                )
    return best_tailored, best_fit, history


# ---------------------------------------------------------------------------
# Step 2 — DOCX renderer (modern design)
# ---------------------------------------------------------------------------

ACCENT = "2B4C7E"  # dark steel blue
DARK = "1A1A1A"    # near-black for body text
GRAY = "555555"    # secondary text


def set_spacing(paragraph, before=0, after=0, line=240):
    """Set paragraph spacing in twips."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def set_font(run, size, bold=False, color=DARK, name="Calibri"):
    """Apply font styling."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def setup_page(doc):
    """Configure US Letter, balanced margins."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def setup_bullet_numbering(doc):
    """Create a bullet list numbering definition and return the numId."""
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    abstract_num_id = "1"
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_num_id)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2022")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    pPr.append(ind)
    lvl.append(pPr)

    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Symbol")
    rFonts.set(qn("w:hAnsi"), "Symbol")
    rFonts.set(qn("w:hint"), "default")
    rPr.append(rFonts)
    lvl.append(rPr)

    abstract_num.append(lvl)
    numbering_elm.append(abstract_num)

    num_id = "1"
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_num_id)
    num.append(abstract_ref)
    numbering_elm.append(num)

    return num_id


def add_bullet(doc, text, num_id):
    """Add a bullet paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 10.5, color=DARK)

    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numPr.append(ilvl)
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), num_id)
    numPr.append(numId)
    pPr.append(numPr)

    set_spacing(p, before=0, after=5)
    return p


def add_section_header(doc, text):
    """Add section header with accent-colored bottom border."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    set_font(run, 11.5, bold=True, color=ACCENT)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), ACCENT)
    pBdr.append(bottom)
    pPr.append(pBdr)

    set_spacing(p, before=120, after=40)


def add_header(doc, meta):
    """Render name, title, and contact details."""
    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta["name"].upper())
    set_font(run, 20, bold=True, color=ACCENT)
    run.font.character_spacing = Pt(1.5)
    set_spacing(p, before=0, after=0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta["title"])
    set_font(run, 11.5, color=GRAY)
    set_spacing(p, before=20, after=20)

    # Contact line
    contact_parts = []
    for key in ("location", "phone", "email", "linkedin", "github", "relocation"):
        if meta.get(key):
            contact_parts.append(meta[key])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep = "  \u00b7  "  # middle dot separator
    run = p.add_run(sep.join(contact_parts))
    set_font(run, 8.5, color=GRAY)
    set_spacing(p, before=0, after=60)


def add_summary(doc, summary_text):
    """Render the summary paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(summary_text)
    set_font(run, 10.5, color=DARK)
    set_spacing(p, before=0, after=40)


def add_experience(doc, experience, num_id):
    """Render experience entries.

    Layout (2 lines of meta per entry):
      Line 1: **Title**  |  Company
      Line 2: Location  |  Start – End

    Keeping location off line 1 means long combined/merged titles don't
    push the location into an awkward wrap.
    """
    for job in experience:
        # Line 1: Title + Company
        p = doc.add_paragraph()
        run = p.add_run(job["title"])
        set_font(run, 11, bold=True, color=DARK)
        run = p.add_run(f"  |  {job['company']}")
        set_font(run, 10.5, color=GRAY)
        set_spacing(p, before=60, after=0)

        # Line 2: Location + Dates
        p = doc.add_paragraph()
        run = p.add_run(
            f"{job['location']}  |  {job['start']} \u2013 {job['end']}"
        )
        set_font(run, 10, color=GRAY)
        set_spacing(p, before=0, after=20)

        for bullet in job["bullets"]:
            add_bullet(doc, bullet, num_id)


def add_skills(doc, skills):
    """Render skills as labeled inline lists."""
    for label, items in skills.items():
        p = doc.add_paragraph()
        run = p.add_run(f"{label.title()}: ")
        set_font(run, 10.5, bold=True, color=ACCENT)
        run = p.add_run(", ".join(items))
        set_font(run, 10.5, color=DARK)
        set_spacing(p, before=0, after=20)


def add_projects(doc, projects, num_id):
    """Render projects."""
    for proj in projects:
        p = doc.add_paragraph()
        run = p.add_run(proj["name"])
        set_font(run, 11, bold=True, color=DARK)
        run = p.add_run(f"  |  {', '.join(proj['tech'])}")
        set_font(run, 10.5, color=GRAY)
        set_spacing(p, before=60, after=20)

        add_bullet(doc, proj["description"], num_id)


def add_education(doc, education):
    """Render education entries."""
    for edu in education:
        p = doc.add_paragraph()
        run = p.add_run(edu["degree"])
        set_font(run, 11, bold=True, color=DARK)
        if edu.get("honors"):
            run = p.add_run(f"  |  {edu['honors']}")
            set_font(run, 10.5, color=ACCENT)
        set_spacing(p, before=60, after=0)

        p = doc.add_paragraph()
        run = p.add_run(f"{edu['institution']}  |  {edu['location']}  |  {edu['start']} \u2013 {edu['end']}")
        set_font(run, 10.5, color=GRAY)
        set_spacing(p, before=0, after=0)


def render_docx(data: dict, output_path: Path):
    """Build a one-page ATS-friendly DOCX from tailored JSON."""
    doc = Document()
    setup_page(doc)

    # Initialize bullet numbering
    doc.add_paragraph("", style="List Bullet")
    doc._body._body.remove(doc.paragraphs[-1]._p)
    num_id = setup_bullet_numbering(doc)

    add_header(doc, data["meta"])
    add_summary(doc, data["summary"])

    add_section_header(doc, "Experience")
    add_experience(doc, data["experience"], num_id)

    add_section_header(doc, "Skills")
    add_skills(doc, data["skills"])

    add_section_header(doc, "Projects")
    add_projects(doc, data["projects"], num_id)

    add_section_header(doc, "Education")
    add_education(doc, data["education"])

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def slugify(text: str) -> str:
    """Lowercase, replace spaces/special chars with underscores."""
    text = text.lower().strip()
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    logging.basicConfig(level=logging.INFO, format="%(message)s")
    parser = argparse.ArgumentParser(description="Tailor a resume for a job posting")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--job-text", help="Job posting as raw text")
    group.add_argument("--job-file", help="Path to a .txt file containing the job posting")
    parser.add_argument("--dry-run", action="store_true", help="Print tailored JSON without rendering DOCX")
    parser.add_argument("--max-iterations", type=int, default=3, help="Max fit/aesthetics retry iterations")
    args = parser.parse_args()

    if args.job_file:
        job_description = Path(args.job_file).read_text()
    else:
        job_description = args.job_text
    log.info(f"Job description: {len(job_description)} chars")

    master = json.loads(MASTER_RESUME.read_text())

    if args.dry_run:
        tailored = call_claude(job_description, master)
        company = slugify(tailored["company"])
        job_title = slugify(tailored["job_title"])
        out_dir = SCRIPT_DIR / "output" / company / job_title
        out_dir.mkdir(parents=True, exist_ok=True)
        (out_dir / "job_description.txt").write_text(job_description)
        (out_dir / "tailored_resume.json").write_text(json.dumps(tailored, indent=2))
        print(json.dumps(tailored, indent=2))
        print(f"\nJSON saved to: {out_dir}/tailored_resume.json")
        return

    # Stage into a temp dir so we don't know company/job_title until after iter 0.
    import tempfile
    with tempfile.TemporaryDirectory(prefix="resume_stage_") as staging:
        staging_dir = Path(staging)
        tailored, fit, history = tailor_with_loop(
            job_description, master, staging_dir, max_iterations=args.max_iterations
        )
        company = slugify(tailored["company"])
        job_title = slugify(tailored["job_title"])
        out_dir = SCRIPT_DIR / "output" / company / job_title
        out_dir.mkdir(parents=True, exist_ok=True)
        (out_dir / "job_description.txt").write_text(job_description)
        # Move rendered artifacts from staging to final out_dir
        for f in staging_dir.iterdir():
            shutil.copy2(f, out_dir / f.name)
    (out_dir / "tailored_resume.json").write_text(json.dumps(tailored, indent=2))
    (out_dir / "fit_history.json").write_text(json.dumps(history, indent=2))

    log.info(
        f"Done: {out_dir}/ (pages={fit['pages']} fill={fit['fill_ratio']})"
    )


if __name__ == "__main__":
    main()
